"""Read text content from .docx files."""

from typing import List, Tuple

import docx


def read_docx(path: str) -> str:
    """Return the full text of a .docx file as a single string.

    Text is extracted from paragraphs and table cells, preserving the
    reading order as closely as possible.

    Args:
        path: Filesystem path to the .docx file.

    Returns:
        All document text joined by newlines.
    """
    doc = docx.Document(path)
    parts: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    parts.append(text)

    return "\n".join(parts)


def read_docx_paragraphs(path: str) -> List[Tuple[str, str]]:
    """Return paragraphs from a .docx file together with their style name.

    Args:
        path: Filesystem path to the .docx file.

    Returns:
        List of (paragraph_text, style_name) tuples (empty paragraphs omitted).
    """
    doc = docx.Document(path)
    paragraphs: List[Tuple[str, str]] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append((text, para.style.name))

    return paragraphs
